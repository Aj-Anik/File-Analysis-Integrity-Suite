"""
algorithms/plagiarism.py

High-level orchestration for the Plagiarism Detector module.

Responsibilities:
    * Read raw text out of .txt, .docx, and .pdf files.
    * Run the similarity algorithms (algorithms/similarity.py) over two
      texts.
    * Assemble a single structured result dictionary that the UI and the
      PDF-report exporter both consume.

External libraries used (install via requirements.txt):
    * python-docx  -> reading .docx files
    * pypdf        -> reading .pdf files
"""

import os
import time
from dataclasses import dataclass, field
from typing import Dict, List, Optional

from algorithms.similarity import (
    jaccard_similarity,
    cosine_similarity,
    find_matching_phrases_rabin_karp,
    tokenize,
)


class UnsupportedFileTypeError(Exception):
    """Raised when a file extension is not one of .txt, .docx, .pdf."""
    pass


@dataclass
class PlagiarismResult:
    """Structured result of comparing Text A and Text B."""
    jaccard_score: float
    cosine_score: float
    common_words: List[str]
    unique_to_a: List[str]
    unique_to_b: List[str]
    matching_phrases: List[str]
    word_count_a: int
    word_count_b: int
    matching_word_count: int
    time_taken_seconds: float
    label_a: str = "Text A"
    label_b: str = "Text B"

    @property
    def overall_similarity_percent(self) -> float:
        """
        Blend Jaccard and Cosine into a single headline percentage.
        We average the two so the result isn't overly dominated by either
        a purely set-based or purely frequency-based view of similarity.
        """
        return round(((self.jaccard_score + self.cosine_score) / 2.0) * 100, 2)

    def to_dict(self) -> Dict:
        return {
            "jaccard_score": round(self.jaccard_score * 100, 2),
            "cosine_score": round(self.cosine_score * 100, 2),
            "overall_similarity_percent": self.overall_similarity_percent,
            "common_words": self.common_words,
            "unique_to_a": self.unique_to_a,
            "unique_to_b": self.unique_to_b,
            "matching_phrases": self.matching_phrases,
            "word_count_a": self.word_count_a,
            "word_count_b": self.word_count_b,
            "matching_word_count": self.matching_word_count,
            "time_taken_seconds": self.time_taken_seconds,
            "label_a": self.label_a,
            "label_b": self.label_b,
        }


# ---------------------------------------------------------------------- #
# File reading
# ---------------------------------------------------------------------- #
def read_text_file(path: str) -> str:
    """Read a plain .txt file, trying utf-8 first then falling back gracefully."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="latin-1") as f:
            return f.read()


def read_docx_file(path: str) -> str:
    """Extract all paragraph text from a .docx file using python-docx."""
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required to read .docx files. Install it with: pip install python-docx"
        ) from e

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    # Also pull text out of any tables, since reports/essays sometimes include them
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def read_pdf_file(path: str) -> str:
    """Extract all text from a .pdf file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "pypdf is required to read .pdf files. Install it with: pip install pypdf"
        ) from e

    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    return "\n".join(parts)


def read_any_supported_file(path: str) -> str:
    """
    Dispatch to the correct reader based on file extension.
    Raises UnsupportedFileTypeError for anything other than
    .txt, .docx, .pdf.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return read_text_file(path)
    elif ext == ".docx":
        return read_docx_file(path)
    elif ext == ".pdf":
        return read_pdf_file(path)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported types: .txt, .docx, .pdf"
        )


# ---------------------------------------------------------------------- #
# Main comparison entry point
# ---------------------------------------------------------------------- #
def compare_texts(
    text_a: str,
    text_b: str,
    label_a: str = "Text A",
    label_b: str = "Text B",
    phrase_len: int = 5,
) -> PlagiarismResult:
    """
    Run the full plagiarism-detection pipeline on two raw text strings and
    return a PlagiarismResult bundling every metric the UI needs to display.
    """
    start_time = time.perf_counter()

    jaccard_score, common, unique_a, unique_b = jaccard_similarity(text_a, text_b)
    cosine_score = cosine_similarity(text_a, text_b)
    matching_phrases = find_matching_phrases_rabin_karp(text_a, text_b, phrase_len=phrase_len)

    words_a = tokenize(text_a)
    words_b = tokenize(text_b)

    elapsed = time.perf_counter() - start_time

    return PlagiarismResult(
        jaccard_score=jaccard_score,
        cosine_score=cosine_score,
        common_words=sorted(common),
        unique_to_a=sorted(unique_a),
        unique_to_b=sorted(unique_b),
        matching_phrases=matching_phrases,
        word_count_a=len(words_a),
        word_count_b=len(words_b),
        matching_word_count=len(common),
        time_taken_seconds=round(elapsed, 4),
        label_a=label_a,
        label_b=label_b,
    )


def compare_files(path_a: str, path_b: str, phrase_len: int = 5) -> PlagiarismResult:
    """Convenience wrapper: read two files from disk, then compare_texts() them."""
    text_a = read_any_supported_file(path_a)
    text_b = read_any_supported_file(path_b)
    return compare_texts(
        text_a, text_b,
        label_a=os.path.basename(path_a),
        label_b=os.path.basename(path_b),
        phrase_len=phrase_len,
    )
